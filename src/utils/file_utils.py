import shutil
from fastapi import UploadFile, HTTPException
from pathlib import Path
import uuid
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

TEMP_DIR = Path(tempfile.gettempdir())

def save_upload_file(upload_file: UploadFile) -> str:
    """Save the uploaded file to a temporary location and return the path"""
    try:
        file_extension = Path(upload_file.filename).suffix
        if file_extension.lower() != '.docx':
            raise HTTPException(status_code=400, detail="Only .docx files are allowed")
        temp_file_path = TEMP_DIR / f"{uuid.uuid4()}{file_extension}"
        with temp_file_path.open("wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
        return str(temp_file_path)
    finally:
        upload_file.file.close()

def create_script_docx(script_content: str, topic: str) -> str:
    """Create a DOCX file with the generated script and return its path"""
    doc = Document()
    title = doc.add_heading(topic, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(script_content)
    file_path = TEMP_DIR / f"script_{uuid.uuid4()}.docx"
    doc.save(file_path)
    return str(file_path) 