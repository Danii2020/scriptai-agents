from datetime import datetime
import shutil
from fastapi import FastAPI, HTTPException, BackgroundTasks, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List
import asyncio
import uuid
from enum import Enum
import os
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .langgraph_workflow import run_youtube_script_workflow

# Load environment variables
load_dotenv()

# Get allowed origins from environment variable or use default
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "*").split(",")

# Create temp directory if it doesn't exist
TEMP_DIR = Path("temp_uploads")
TEMP_DIR.mkdir(exist_ok=True)

app = FastAPI(
    title="YouTube Script Generator API",
    description="API for generating YouTube scripts using CrewAI",
    version="1.0.0"
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=3600,  # Cache preflight requests for 1 hour
)

# In-memory task store
tasks = {}

class TaskStatus(str, Enum):
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"

class ScriptResponse(BaseModel):
    task_id: str
    status: TaskStatus
    result: Optional[str] = None
    error: Optional[str] = None
    file_path: Optional[str] = None

async def save_upload_file(upload_file: UploadFile) -> str:
    """Save the uploaded file to a temporary location and return the path"""
    try:
        # Create a unique filename
        file_extension = Path(upload_file.filename).suffix
        if file_extension.lower() != '.docx':
            raise HTTPException(status_code=400, detail="Only .docx files are allowed")
            
        temp_file_path = TEMP_DIR / f"{uuid.uuid4()}{file_extension}"
        
        # Save the file
        with temp_file_path.open("wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
            
        return str(temp_file_path)
    finally:
        upload_file.file.close()

def create_script_docx(script_content: str, topic: str) -> str:
    """Create a DOCX file with the generated script and return its path"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(topic, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add script content
    doc.add_paragraph(script_content)
    
    # Create temporary file
    temp_dir = Path("temp_docs")
    temp_dir.mkdir(exist_ok=True)
    file_path = temp_dir / f"script_{uuid.uuid4()}.docx"
    
    # Save the document
    doc.save(file_path)
    return str(file_path)


async def run_langgraph_task(task_id: str, topic: str, tones: List[str], file_path: Optional[str]):
    try:
        tasks[task_id]["status"] = TaskStatus.RUNNING
        # Run the new LangGraph workflow (sync for now; wrap in thread if needed)
        script_content = await asyncio.to_thread(
            run_youtube_script_workflow,
            topic=topic,
            tones=", ".join(tones),
            file_path=file_path,
            current_year=str(datetime.now().year)
        )
        docx_path = create_script_docx(script_content, topic)
        tasks[task_id].update({
            "status": TaskStatus.COMPLETED,
            "result": script_content,
            "file_path": docx_path
        })
    except Exception as e:
        tasks[task_id].update({
            "status": TaskStatus.FAILED,
            "error": str(e)
        })
    finally:
        if file_path and file_path != "/Users/danielerazo/Documents/yt-scripts/script-template-en.docx":
            try:
                os.remove(file_path)
            except:
                pass

@app.post("/generate-script", response_model=ScriptResponse)
async def generate_script(
    topic: str = Form(...),
    tones: List[str] = Form(["professional"]),
    file_name: Optional[UploadFile] = File(None),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    task_id = str(uuid.uuid4())
    tasks[task_id] = {
        "status": TaskStatus.PENDING,
        "result": None,
        "error": None
    }

    file_path = "/Users/danielerazo/Documents/yt-scripts/script-template-en.docx"
    if file_name:
        file_path = await save_upload_file(file_name)
    background_tasks.add_task(run_langgraph_task, task_id, topic, tones, file_path)

    return ScriptResponse(
        task_id=task_id,
        status=TaskStatus.PENDING
    )

@app.get("/task/{task_id}", response_model=ScriptResponse)
async def get_task_status(task_id: str):
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = tasks[task_id]
    return ScriptResponse(
        task_id=task_id,
        status=task["status"],
        result=task["result"],
        error=task["error"],
        file_path=task.get("file_path")
    )

@app.get("/download-script/{task_id}")
async def download_script(task_id: str):
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = tasks[task_id]
    if task["status"] != TaskStatus.COMPLETED or not task.get("file_path"):
        raise HTTPException(status_code=400, detail="Script not ready for download")
    
    file_path = task["file_path"]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Script file not found")
    
    return FileResponse(
        path=file_path,
        filename=f"script_{task_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/health")
async def health_check():
    return {"status": "healthy"}
